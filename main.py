"""Herf — InPage to Unicode converter for Windows.

Run:  python main.py
All processing is local. No file content leaves this machine.
"""

import json
import os
import sys
import webview

from herf.engine import convert_bytes, convert, decode_bytes

APP_DIR = os.path.dirname(os.path.abspath(__file__))


class Api:
    def __init__(self):
        self.user_mapping = {}
        self.last_text = ""
        # Per-line styles for the last .inp conversion, aligned with
        # last_text.split("\n"). Each entry is {"size": pt|None,
        # "bold": bool} or None (nothing recovered for that line).
        # None as a whole means the current text has no style data
        # (pasted text, .txt files) — export falls back to flat 14pt.
        self.last_styles = None

    # ---------- conversion ----------

    def convert_pasted(self, pasted, opts):
        """Pasted text arrives already unicode-decoded by the browser;
        re-encode via cp1252 to recover original bytes where possible."""
        try:
            raw = pasted.encode("cp1252", errors="replace")
        except Exception:
            raw = pasted.encode("utf-8")
        return self._run(raw, opts)

    def convert_file(self, path, opts):
        from herf.inp_reader import is_inp, extract_document
        if path.lower().endswith(".inp") or is_inp(path):
            blocks, block_styles = extract_document(path)
            kw = dict(user_mapping=self.user_mapping,
                      latin_digits=opts.get("latinDigits", False),
                      strip_kashida=opts.get("stripKashida", False))
            # Convert paragraph-by-paragraph so each output line keeps
            # its recovered style. Verified byte-identical to whole-text
            # conversion on the reference file (rules never span a
            # paragraph mark). Reports are merged across paragraphs.
            lines, styles, merged = [], [], {}
            for bi, blk in enumerate(blocks):
                if bi:                       # blocks joined by a blank line
                    lines.append("")
                    styles.append(None)
                paras = blk.split("\n")
                pstyles = block_styles[bi]
                for j, para in enumerate(paras):
                    t, rep = convert(para, **kw)
                    st = pstyles[j] if j < len(pstyles) else None
                    lines.append(t)
                    styles.append({"size": st["size"], "bold": st["bold"]}
                                  if st else None)
                    for e in rep:
                        m = merged.setdefault(e["code"], e)
                        if m is not e:
                            m["count"] += e["count"]
                            m["contexts"] = (m["contexts"] + e["contexts"])[:5]
            # whole-text convert() strips leading/trailing blank lines;
            # mirror that so behavior and alignment stay identical
            while lines and not lines[0].strip():
                lines.pop(0); styles.pop(0)
            while lines and not lines[-1].strip():
                lines.pop(); styles.pop()
            text = "\n".join(lines)
            report = sorted(merged.values(), key=lambda e: -e["count"])
            self.last_text = text
            self.last_styles = styles
            return {"text": text, "report": report, "chars": len(text),
                    "unmapped": sum(e["count"] for e in report),
                    "blocks": len(blocks), "source": "inp",
                    "styled": sum(1 for s in styles if s)}
        with open(path, "rb") as f:
            raw = f.read()
        return self._run(raw, opts)

    def _run(self, raw, opts):
        text, report = convert_bytes(
            raw,
            user_mapping=self.user_mapping,
            latin_digits=opts.get("latinDigits", False),
            strip_kashida=opts.get("stripKashida", False),
        )
        self.last_text = text
        self.last_styles = None      # pasted/.txt input carries no style data
        return {"text": text, "report": report,
                "chars": len(text), "unmapped": sum(e["count"] for e in report)}


    # ---------- cleanup and Quran check ----------

    def fix_spelling(self, text):
        from herf.textfix import fix_spelling
        fixed, changes = fix_spelling(text)
        self.last_text = fixed
        return {"text": fixed, "changes": changes}

    def quran_scan(self, text):
        from herf.quran import QuranIndex, scan
        if not hasattr(self, "_qindex"):
            self._qindex = QuranIndex()
        matches = scan(text, self._qindex)
        return {"matches": matches, "count": len(matches)}

    def set_text(self, text):
        """UI pushes its authoritative text after client-side edits."""
        self.last_text = text
        return {"ok": True}

    # ---------- custom mappings ----------

    def add_mapping(self, code_hex, replacement):
        """code_hex like '0x04 0xE2'. Store and return updated mapping list."""
        parts = code_hex.split()
        b = int(parts[1], 16)
        undefined = {0x81, 0x8D, 0x8F, 0x90, 0x9D}
        ch = chr(b) if b in undefined else bytes([b]).decode("cp1252")
        self.user_mapping["\u0004" + ch] = replacement
        return {"ok": True, "count": len(self.user_mapping)}

    # ---------- files ----------

    def pick_files(self):
        w = webview.windows[0]
        result = w.create_file_dialog(
            webview.OPEN_DIALOG, allow_multiple=True,
            file_types=("InPage files (*.inp;*.INP;*.txt;*.TXT)",
                        "All files (*.*)"))
        return list(result) if result else []

    def pick_folder(self):
        w = webview.windows[0]
        result = w.create_file_dialog(webview.FOLDER_DIALOG)
        if not result:
            return []
        folder = result[0]
        out = []
        for name in sorted(os.listdir(folder)):
            if name.lower().endswith((".txt", ".inp")):
                out.append(os.path.join(folder, name))
        return out

    # ---------- export ----------

    def save_txt(self, text):
        w = webview.windows[0]
        path = w.create_file_dialog(
            webview.SAVE_DIALOG, save_filename="converted.txt")
        if not path:
            return {"ok": False}
        if isinstance(path, (list, tuple)):
            path = path[0]
        with open(path, "w", encoding="utf-8") as f:
            f.write(text)
        return {"ok": True, "path": path}

    def save_docx(self, text, font_name):
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            return {"ok": False,
                    "error": "python-docx is not installed. Run: pip install python-docx"}

        w = webview.windows[0]
        path = w.create_file_dialog(
            webview.SAVE_DIALOG, save_filename="converted.docx")
        if not path:
            return {"ok": False}
        if isinstance(path, (list, tuple)):
            path = path[0]

        lines = text.split("\n")
        # Recovered .inp styles apply only while they still line up with
        # the text being saved; after freehand edits they are dropped
        # rather than mis-applied.
        styles = self.last_styles
        if not styles or len(styles) != len(lines):
            styles = None

        doc = Document()
        for idx, para_text in enumerate(lines):
            st = styles[idx] if styles else None
            if not para_text.strip() and styles is None:
                # legacy flat export: skip blanks (no styling to keep)
                continue
            size = st["size"] if st and st["size"] else 14
            bold = bool(st and st["bold"])
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # RTL paragraph direction
            pPr = p._p.get_or_add_pPr()
            bidi = OxmlElement("w:bidi")
            pPr.append(bidi)
            if not para_text.strip():
                # blank spacer paragraph from the source: keep it, and
                # size the paragraph mark so the gap height matches
                mark_rPr = OxmlElement("w:rPr")
                for tag in ("w:sz", "w:szCs"):
                    el = OxmlElement(tag)
                    el.set(qn("w:val"), str(int(size * 2)))
                    mark_rPr.append(el)
                pPr.insert(0, mark_rPr)
                continue
            run = p.add_run(para_text)
            run.font.name = font_name
            run.font.size = Pt(size)
            if bold:
                run.font.bold = True
            rPr = run._r.get_or_add_rPr()
            rtl = OxmlElement("w:rtl")
            rPr.append(rtl)
            cs = OxmlElement("w:szCs")
            cs.set(qn("w:val"), str(int(size * 2)))
            rPr.append(cs)
            if bold:                        # bold for complex scripts too
                bCs = OxmlElement("w:bCs")
                rPr.append(bCs)
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:cs"), font_name)
        doc.save(path)
        return {"ok": True, "path": path}

    def list_fonts(self):
        """Return installed font family names on Windows; fallback list otherwise."""
        try:
            import winreg
            fonts = set()
            key = winreg.OpenKey(
                winreg.HKEY_LOCAL_MACHINE,
                r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts")
            i = 0
            while True:
                try:
                    name, _, _ = winreg.EnumValue(key, i)
                    fonts.add(name.split(" (")[0])
                    i += 1
                except OSError:
                    break
            return sorted(fonts)
        except Exception:
            return ["Jameel Noori Nastaleeq", "Urdu Typesetting",
                    "Noto Nastaliq Urdu", "Arial"]


def main():
    api = Api()
    webview.create_window(
        "Herf",
        os.path.join(APP_DIR, "herf", "ui", "index.html"),
        js_api=api,
        width=1280, height=820, min_size=(980, 640),
        background_color="#F7F1EA",
    )
    webview.start()


if __name__ == "__main__":
    main()
